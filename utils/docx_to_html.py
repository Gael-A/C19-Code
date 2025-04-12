import os
import zipfile
import base64
from docx import Document

# --- Funciones de extracción de imágenes ---
def extract_images(docx_path):
    """
    Extrae las imágenes del archivo DOCX (ubicadas en word/media) y
    las devuelve en un diccionario donde la clave es el nombre del archivo (ruta relativa)
    y el valor es el string de la imagen en formato data URI.
    """
    images = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for file in z.namelist():
            if file.startswith("word/media/"):
                image_data = z.read(file)
                ext = os.path.splitext(file)[1][1:].lower()  # extensión sin el punto
                # Determinar el tipo MIME según la extensión
                if ext in ["jpg", "jpeg"]:
                    mime = "image/jpeg"
                elif ext == "png":
                    mime = "image/png"
                elif ext == "gif":
                    mime = "image/gif"
                else:
                    mime = "image/" + ext
                b64 = base64.b64encode(image_data).decode("utf-8")
                images[file] = f"data:{mime};base64,{b64}"
    return images

def get_image_from_run(run, images, document):
    """
    Busca en un run (posible imagen) elementos de dibujo (w:drawing) para obtener el r:embed,
    luego busca en las partes relacionadas la imagen correspondiente y la devuelve en formato <img>.
    """
    drawing_elements = run._element.xpath(".//w:drawing")
    if drawing_elements:
        blips = drawing_elements[0].xpath(".//a:blip")
        if blips:
            # Extraer el atributo embed que contiene el id de la imagen
            embed = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            # Obtener la parte relacionada (la imagen) y su nombre de archivo
            if embed in document.part._rels:
                rel = document.part._rels[embed]
                target = rel._target  # target es un objeto ImagePart
                # Usamos target.partname para obtener el string de la ruta
                image_path = target.partname.lstrip(
                    "/"
                )  # por ejemplo: "word/media/image1.png"
                if image_path in images:
                    return f'<img src="{images[image_path]}">'
    return ""


# --- Funciones de conversión a HTML ---
def convert_run(run, document):
    """
    Convierte un run (fragmento de texto con formato) a HTML, aplicando etiquetas
    según si está en negrita, cursiva, subrayado, tachado o con tamaño de fuente.
    """
    text = run.text
    if not text:
        return ""
    # Aplicar formatos: se anidan las etiquetas según el formato detectado.
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    # Para tachado, python-docx lo almacena en run.font.strike
    if run.font.strike:
        text = f"<strike>{text}</strike>"
    # Si se ha definido tamaño de fuente, se agrega un span con estilo
    if run.font.size:
        # El tamaño se obtiene en puntos (pt)
        size_pt = run.font.size.pt
        text = f'<span style="font-size:{size_pt}pt;">{text}</span>'
    return text

def convert_paragraph(paragraph, images, document):
    # Eliminar el manejo de listas aquí (se procesarán aparte)
    img_html = "".join(
        [get_image_from_run(run, images, document) for run in paragraph.runs]
    )

    if not paragraph.text.strip() and not img_html:
        return "<br>\n"

    align = paragraph.alignment
    align_class = {
        0: "align-left",
        1: "align-center",
        2: "align-right",
        3: "align-justify",
    }.get(align, "align-left")

    if "Heading" in paragraph.style.name:
        try:
            level = int(paragraph.style.name.split(" ")[-1])
        except:
            level = 1
        inner = "".join([convert_run(run, document) for run in paragraph.runs])
        return f"<h{level}>{inner}</h{level}>\n"

    inner = "".join([convert_run(run, document) for run in paragraph.runs])
    return f'<p class="{align_class}">{inner}{img_html}</p>\n'

def is_list_paragraph(paragraph):
    # Mejorar la detección de listas
    style_name = paragraph.style.name.lower()
    return any(keyword in style_name for keyword in ["list", "bullet", "number"])

def convert_list_items(list_paragraphs, images, document):
    items = []
    for para in list_paragraphs:
        p_html = convert_paragraph(para, images, document).strip()
        if p_html.startswith("<p"):
            start = p_html.find(">") + 1
            end = p_html.rfind("</p>")
            inner = p_html[start:end]
        else:
            inner = p_html
        items.append(f"<li>{inner}</li>")

    # Determinar si es lista ordenada
    first_style = list_paragraphs[0].style.name.lower()
    list_tag = "ol" if "number" in first_style else "ul"
    return f"<{list_tag}>\n" + "\n".join(items) + f"\n</{list_tag}>\n"

def convert_table(table, images, document):
    """
    Convierte una tabla DOCX a HTML, procesando sus filas y celdas.
    """
    html = "<table border='1'>\n"
    for row in table.rows:
        html += "  <tr>\n"
        for cell in row.cells:
            cell_content = ""
            for paragraph in cell.paragraphs:
                cell_content += convert_paragraph(paragraph, images, document)
            html += f"    <td>{cell_content}</td>\n"
        html += "  </tr>\n"
    html += "</table>\n"
    return html

def convert_docx_to_html(docx_path):
    images = extract_images(docx_path)
    document = Document(docx_path)

    html_content = """"""

    elements = list(document.element.body)
    i = 0
    while i < len(elements):
        element = elements[i]
        tag = element.tag

        if tag.endswith("p"):
            para = next((p for p in document.paragraphs if p._element == element), None)
            if para and is_list_paragraph(para):
                # Agrupar párrafos consecutivos de lista
                list_paras = []
                while i < len(elements) and elements[i].tag.endswith("p"):
                    current_para = next(
                        (p for p in document.paragraphs if p._element == elements[i]),
                        None,
                    )
                    if current_para and is_list_paragraph(current_para):
                        list_paras.append(current_para)
                        i += 1
                    else:
                        break
                html_content += convert_list_items(list_paras, images, document) + "\n"
            else:
                html_content += convert_paragraph(para, images, document)
                i += 1
        elif tag.endswith("tbl"):
            table = next((t for t in document.tables if t._element == element), None)
            if table:
                html_content += convert_table(table, images, document) + "\n"
            i += 1
        else:
            i += 1

    return html_content

# convert_docx_to_html("input.docx", "salida.html")
